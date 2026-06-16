import os
from typing import Optional

def extract_text(file_path: str, file_type: str) -> Optional[str]:
    try:
        if file_type == "pdf":
            return extract_from_pdf(file_path)
        elif file_type in ["docx", "doc"]:
            return extract_from_docx(file_path)
    except Exception as e:
        print(f"Text extraction error: {e}")
        return None

def extract_from_pdf(file_path: str) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except ImportError:
        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"PDF extraction failed: {e}")

def extract_from_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {e}")
